"""ACE Export Package Builder.

Builds a ZIP file containing the full 10-file export package:
    ACE-change-export-EXP-YYYY-NNNN/
    ├── Read-Me.txt
    ├── Summary.xlsx
    ├── Change-Log.xlsx
    ├── Summary.docx
    ├── Summary.pdf
    ├── Change-Log.docx
    ├── Change-Log.pdf
    ├── Changes.csv
    ├── Changes.json
    └── manifest.json

Uses Australian English throughout.
Applies AuditCo Word rules: Capital Case headings, no text boxes, editable.
"""

from __future__ import annotations

import csv
import json
import os
import shutil
import subprocess
import tempfile
import zipfile
from datetime import UTC, datetime
from hashlib import sha256
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Side, Border
from openpyxl.utils import get_column_letter

from src.ace.workbench.change_record import ChangeRecord
from src.ace.workbench.notion_publisher import PublicationResult

# ── Australian English constant strings ────────────────────────

NOT_AUTHORITATIVE_NOTICE = (
    "This document is not authoritative. "
    "It is generated from fictional workbench data for demonstration purposes only."
)
WRITEBACK_NOT_PERMITTED_NOTICE = (
    "Writeback is not permitted. "
    "Changes must originate in the ACE workbench, not in this export."
)
FICTIONAL_DATA_NOTICE = (
    "All records in this export are fictional. "
    "No real client data, photos, or secrets are included."
)


def _utc_now_compact() -> str:
    """Return compact UTC timestamp: YYYYMMDDTHHMMSSZ"""
    return datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")


def _auditco_heading(text: str) -> str:
    """Return the heading in Capital Case (AuditCo Word rule).

    Uses manual capitalisation to preserve known acronyms (ACE, CHG,
    SNP, EXP, ID, CSV, JSON, PDF, XLSX, DOCX) that str.title() would
    incorrectly lowercase.
    """
    _PRESERVED = frozenset({
        "ACE", "CHG", "SNP", "EXP", "ID", "CSV", "JSON", "PDF",
        "XLSX", "DOCX", "URL", "API", "SHA", "N/A",
    })

    def _capitalise_word(w: str) -> str:
        upper = w.upper()
        if upper in _PRESERVED:
            return upper
        # Handle change IDs like CHG-3EC48B35D024
        if "-" in w:
            parts = w.split("-", 1)
            prefix = parts[0].upper()
            if prefix in _PRESERVED:
                return f"{prefix}-{parts[1].upper()}"
        return w[0].upper() + w[1:].lower() if w else w

    return " ".join(_capitalise_word(w) for w in text.split())


# ── Export ID generation ───────────────────────────────────────

def generate_export_id(engagement_id: str, idempotency_key: str = "") -> str:
    """Generate a time-based export ID.

    Format: EXP-YYYY-NNNNNNNNNNNN where YYYY is year and NNNNNNNNNNNN is a
    12-digit hex fragment from a SHA-256 hash of engagement_id + timestamp
    + idempotency_key.  Including the idempotency key guarantees
    uniqueness even when two requests arrive during the same second.
    """
    year = str(datetime.now(UTC).year)
    payload = f"{engagement_id}|{_utc_now_compact()}|{idempotency_key}".encode()
    hex_frag = sha256(payload).hexdigest()[:12].upper()
    return f"EXP-{year}-{hex_frag}"


# ── Formula injection guard ────────────────────────────────────

def _xlsx_safe(value: object) -> str:
    """Escape formula-injection prefixes in spreadsheet cell values.

    Prepends a single-quote prefix to values that start with ``=``,
    ``+``, ``-``, or ``@`` so spreadsheet applications treat them as
    literal text rather than executable formulas.
    """
    s = str(value) if value is not None else ""
    if s and s[0] in ("=", "+", "-", "@"):
        return "'" + s
    return s


# ── Summary XLSX ───────────────────────────────────────────────

def build_summary_xlsx(
    export_id: str,
    snapshot_id: str,
    export_time: str,
    engagement_id: str,
    engagement_title: str,
    changes: list[ChangeRecord],
    warnings: list[dict[str, object]],
    excluded_count: int,
    publication_result: PublicationResult | None,
) -> BytesIO:
    """Build the Summary.xlsx workbook."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"

    header_font = Font(bold=True, size=12)
    header_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    def _write_row(row_num: int, label: str, value: str) -> None:
        c1 = ws.cell(row=row_num, column=1, value=label)
        c2 = ws.cell(row=row_num, column=2, value=_xlsx_safe(value))
        c1.font = Font(bold=True)
        for c in (c1, c2):
            c.border = thin_border

    # Header
    ws.merge_cells("A1:B1")
    h = ws.cell(row=1, column=1, value=_auditco_heading("ACE Change Export Summary"))
    h.font = Font(bold=True, size=14)
    h.alignment = Alignment(horizontal="center")

    row = 3
    _write_row(row, "Export ID", export_id)
    row += 1
    _write_row(row, "Snapshot ID", snapshot_id)
    row += 1
    _write_row(row, "Export Time", export_time)
    row += 1
    _write_row(row, "Engagement Reference", engagement_id)
    row += 1
    _write_row(row, "Engagement Title", engagement_title)
    row += 2

    # Counts by record type
    counts: dict[str, int] = {}
    for ch in changes:
        counts[ch.record_type] = counts.get(ch.record_type, 0) + 1

    ws.cell(row=row, column=1, value="Change Counts By Record Type").font = header_font
    ws.cell(row=row, column=2, value="").font = header_font
    row += 1
    for rtype, cnt in sorted(counts.items()):
        _write_row(row, rtype, str(cnt))
        row += 1

    row += 1
    _write_row(row, "Total Changes", str(len(changes)))
    row += 1
    _write_row(row, "Excluded Records", str(excluded_count))
    row += 2

    # Publication status
    if publication_result is not None:
        _write_row(row, "Publication Status", "Published" if publication_result.published else "Failed")
        row += 1
        _write_row(row, "Publication ID", publication_result.publication_id)
        row += 2

    # Warnings
    if warnings:
        ws.cell(row=row, column=1, value="Warnings").font = header_font
        ws.cell(row=row, column=2, value="").font = header_font
        row += 1
        for w in warnings:
            _write_row(row, str(w.get("level", "")), str(w.get("detail", "")))
            row += 1
        row += 1

    # Linked change IDs
    ws.cell(row=row, column=1, value="Linked Change IDs").font = header_font
    ws.cell(row=row, column=2, value="").font = header_font
    row += 1
    for ch in changes:
        _write_row(row, ch.change_id, f"{ch.record_id} ({ch.change_type})")
        row += 1

    row += 2
    # Notices
    for notice in (NOT_AUTHORITATIVE_NOTICE, WRITEBACK_NOT_PERMITTED_NOTICE, FICTIONAL_DATA_NOTICE):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
        c = ws.cell(row=row, column=1, value=notice)
        c.font = Font(italic=True, size=10, color="666666")
        row += 1

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 60

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


# ── Change-Log XLSX ────────────────────────────────────────────

def build_change_log_xlsx(changes: list[ChangeRecord]) -> BytesIO:
    """Build the Change-Log.xlsx workbook."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Change Log"

    headers = [
        "Change ID", "Export ID", "Record ID", "Snapshot ID",
        "Evidence ID", "Idempotency Key", "Timestamp",
        "Change Type", "Record Type", "Label", "Detail",
    ]
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")

    for col_idx, hdr in enumerate(headers, 1):
        c = ws.cell(row=1, column=col_idx, value=hdr)
        c.font = header_font
        c.fill = header_fill

    for row_idx, ch in enumerate(changes, 2):
        values = [
            ch.change_id, ch.export_id, ch.record_id, ch.snapshot_id,
            ch.evidence_id or "", ch.idempotency_key, ch.timestamp,
            ch.change_type, ch.record_type, _xlsx_safe(ch.label), _xlsx_safe(ch.detail),
        ]
        for col_idx, val in enumerate(values, 1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    for col_idx in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = max(
            16, len(str(ws.cell(row=1, column=col_idx).value or "")) + 4
        )

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


# ── Changes CSV ────────────────────────────────────────────────

def build_changes_csv(changes: list[ChangeRecord]) -> StringIO:
    """Build Changes.csv with all change records."""
    output = StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow([
        "change_id", "export_id", "record_id", "snapshot_id",
        "evidence_id", "idempotency_key", "timestamp",
        "change_type", "record_type", "label", "detail",
    ])
    for ch in changes:
        writer.writerow([
            ch.change_id, ch.export_id, ch.record_id, ch.snapshot_id,
            ch.evidence_id or "", ch.idempotency_key, ch.timestamp,
            ch.change_type, ch.record_type,
            _xlsx_safe(ch.label), _xlsx_safe(ch.detail),
        ])
    output.seek(0)
    return output


# ── Changes JSON ───────────────────────────────────────────────

def build_changes_json(changes: list[ChangeRecord]) -> BytesIO:
    """Build Changes.json with all change records."""
    data = [ch.as_dict() for ch in changes]
    output = BytesIO()
    output.write(json.dumps(data, indent=2).encode("utf-8"))
    output.seek(0)
    return output


# ── DOCX builder ───────────────────────────────────────────────

def _add_auditco_docx_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading in Capital Case per AuditCo Word rules."""
    heading = doc.add_heading(_auditco_heading(text), level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_notice_paragraph(doc: Document, text: str) -> None:
    """Add an italic notice paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = None  # Use default colour
    p.paragraph_format.space_before = Pt(6)


def build_summary_docx(
    export_id: str,
    snapshot_id: str,
    export_time: str,
    engagement_id: str,
    engagement_title: str,
    changes: list[ChangeRecord],
    warnings: list[dict[str, object]],
    excluded_count: int,
    publication_result: PublicationResult | None,
) -> BytesIO:
    """Build the Summary.docx file."""
    doc = Document()

    # Title
    title = doc.add_heading(_auditco_heading("ACE Change Export Summary"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    _add_auditco_docx_heading(doc, "Export Details", level=1)
    meta = [
        ("Export ID", export_id),
        ("Snapshot ID", snapshot_id),
        ("Export Time", export_time),
        ("Engagement Reference", engagement_id),
        ("Engagement Title", engagement_title),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    # Counts
    _add_auditco_docx_heading(doc, "Change Summary", level=1)
    counts: dict[str, int] = {}
    for ch in changes:
        counts[ch.record_type] = counts.get(ch.record_type, 0) + 1
    for rtype, cnt in sorted(counts.items()):
        p = doc.add_paragraph()
        p.add_run(f"{rtype}: ").bold = True
        p.add_run(str(cnt))
    p = doc.add_paragraph()
    p.add_run("Total Changes: ").bold = True
    p.add_run(str(len(changes)))
    p = doc.add_paragraph()
    p.add_run("Excluded Records: ").bold = True
    p.add_run(str(excluded_count))

    # Publication
    if publication_result is not None:
        _add_auditco_docx_heading(doc, "Publication", level=1)
        status = "Published" if publication_result.published else "Failed"
        p = doc.add_paragraph()
        p.add_run("Publication Status: ").bold = True
        p.add_run(status)
        if publication_result.publication_id:
            p2 = doc.add_paragraph()
            p2.add_run("Publication ID: ").bold = True
            p2.add_run(publication_result.publication_id)

    # Warnings
    if warnings:
        _add_auditco_docx_heading(doc, "Warnings", level=1)
        for w in warnings:
            p = doc.add_paragraph()
            p.add_run(f"{w.get('level', '')}: ").bold = True
            p.add_run(str(w.get("detail", "")))

    # Linked Change IDs
    _add_auditco_docx_heading(doc, "Linked Change Identifiers", level=1)
    for ch in changes:
        p = doc.add_paragraph()
        p.add_run(f"{ch.change_id}: {ch.record_id} ({ch.change_type}) — {ch.record_type}")

    # Notices
    doc.add_page_break()
    _add_auditco_docx_heading(doc, "Notices", level=1)
    for notice in (NOT_AUTHORITATIVE_NOTICE, WRITEBACK_NOT_PERMITTED_NOTICE, FICTIONAL_DATA_NOTICE):
        _add_notice_paragraph(doc, notice)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def build_change_log_docx(changes: list[ChangeRecord]) -> BytesIO:
    """Build the Change-Log.docx file."""
    doc = Document()

    title = doc.add_heading(_auditco_heading("ACE Change Log"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_auditco_docx_heading(doc, "Change Records", level=1)

    for ch in changes:
        _add_auditco_docx_heading(doc, ch.change_id, level=2)
        fields = [
            ("Export ID", ch.export_id),
            ("Record ID", ch.record_id),
            ("Snapshot ID", ch.snapshot_id),
            ("Evidence ID", ch.evidence_id or "N/A"),
            ("Idempotency Key", ch.idempotency_key),
            ("Timestamp", ch.timestamp),
            ("Change Type", ch.change_type),
            ("Record Type", ch.record_type),
            ("Label", ch.label),
        ]
        for label, value in fields:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)
        if ch.detail:
            p = doc.add_paragraph()
            p.add_run("Detail: ").bold = True
            p.add_run(ch.detail)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ── DOCX → PDF via LibreOffice ─────────────────────────────────

def docx_to_pdf(docx_bytes: BytesIO, toolchain_doctor_result: bool = True) -> BytesIO:
    """Convert a DOCX BytesIO to PDF using LibreOffice headless.

    Must be called after document-toolchain-doctor -Probe confirms
    LibreOffice is available.
    """
    if not toolchain_doctor_result:
        raise RuntimeError(
            "document-toolchain-doctor -Probe reported missing tools — PDF rendering unavailable"
        )

    candidates = ("libreoffice", "soffice.com") if os.name == "nt" else ("libreoffice",)
    libreoffice = next(
        (shutil.which(candidate) for candidate in candidates if shutil.which(candidate)),
        None,
    )
    if libreoffice is None:
        raise RuntimeError("LibreOffice not found on PATH — PDF rendering unavailable")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        docx_path = tmp / "input.docx"
        docx_path.write_bytes(docx_bytes.getvalue())

        result = subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=30,
            env={**os.environ, "HOME": str(tmp)},
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice PDF conversion failed: {result.stderr.strip()}"
            )

        pdf_path = tmp / "input.pdf"
        if not pdf_path.exists():
            raise RuntimeError("LibreOffice did not produce input.pdf")

        output = BytesIO(pdf_path.read_bytes())
        output.seek(0)
        return output


# ── Read-Me.txt ────────────────────────────────────────────────

def build_readme(
    export_id: str, snapshot_id: str, engagement_id: str,
    export_time: str, skip_pdf: bool = False,
) -> StringIO:
    """Build the Read-Me.txt file, listing only files actually packaged."""
    contents = [
        "  Summary.xlsx      — Export summary with counts, warnings, notices",
        "  Change-Log.xlsx   — All change records in tabular format",
        "  Summary.docx      — Export summary (Word, editable)",
    ]
    if not skip_pdf:
        contents.append("  Summary.pdf       — Export summary (PDF, read-only)")
    contents.extend([
        "  Change-Log.docx   — Change records (Word, editable)",
    ])
    if not skip_pdf:
        contents.append("  Change-Log.pdf    — Change records (PDF, read-only)")
    contents.extend([
        "  Changes.csv       — All change records as CSV",
        "  Changes.json      — All change records as JSON",
        "  manifest.json     — Export manifest with checksums",
    ])

    lines = [
        f"ACE Change Export — {export_id}",
        "=" * 50,
        "",
        f"Export ID:       {export_id}",
        f"Snapshot ID:     {snapshot_id}",
        f"Engagement:      {engagement_id}",
        f"Export Time:     {export_time}",
        "",
        "Contents:",
        *contents,
        "",
        NOT_AUTHORITATIVE_NOTICE,
        WRITEBACK_NOT_PERMITTED_NOTICE,
        FICTIONAL_DATA_NOTICE,
        "",
    ]
    output = StringIO()
    output.write("\n".join(lines))
    output.seek(0)
    return output


# ── Manifest ───────────────────────────────────────────────────

def build_manifest(
    export_id: str,
    snapshot_id: str,
    engagement_id: str,
    file_checksums: dict[str, str],
    change_count: int,
    warning_count: int,
) -> BytesIO:
    """Build manifest.json with file checksums."""
    manifest: dict[str, object] = {
        "export_id": export_id,
        "snapshot_id": snapshot_id,
        "engagement_id": engagement_id,
        "export_time": _utc_now_compact(),
        "total_changes": change_count,
        "total_warnings": warning_count,
        "notices": [
            NOT_AUTHORITATIVE_NOTICE,
            WRITEBACK_NOT_PERMITTED_NOTICE,
            FICTIONAL_DATA_NOTICE,
        ],
        "files": file_checksums,
    }
    output = BytesIO()
    output.write(json.dumps(manifest, indent=2).encode("utf-8"))
    output.seek(0)
    return output


# ── ZIP Package Builder ────────────────────────────────────────

def build_export_zip(
    export_id: str,
    snapshot_id: str,
    export_time: str,
    engagement_id: str,
    engagement_title: str,
    changes: list[ChangeRecord],
    warnings: list[dict[str, object]],
    excluded_count: int,
    publication_result: PublicationResult | None,
    skip_pdf: bool = False,
) -> BytesIO:
    """Build the complete ACE change export ZIP package.

    Returns a BytesIO containing the ZIP file. The ZIP contains all 10
    files in a single directory named ACE-change-export-{export_id}/.
    """
    dirname = f"ACE-change-export-{export_id}"

    # Build all file components
    summary_xlsx = build_summary_xlsx(
        export_id, snapshot_id, export_time,
        engagement_id, engagement_title,
        changes, warnings, excluded_count, publication_result,
    )
    change_log_xlsx = build_change_log_xlsx(changes)
    summary_docx = build_summary_docx(
        export_id, snapshot_id, export_time,
        engagement_id, engagement_title,
        changes, warnings, excluded_count, publication_result,
    )
    change_log_docx = build_change_log_docx(changes)

    # PDF conversion — probe toolchain first, then convert.
    # Fail explicitly when the toolchain is unavailable so callers
    # never receive an incomplete export that silently omits PDFs.
    summary_pdf: BytesIO | None = None
    change_log_pdf: BytesIO | None = None
    if not skip_pdf:
        from src.ace.workbench.document_toolchain_doctor import probe as probe_toolchain

        doctor_report = probe_toolchain()
        if not doctor_report.all_available:
            missing = [r.tool for r in doctor_report.results if not r.available]
            raise RuntimeError(
                f"Document toolchain unavailable — PDF export requires: {', '.join(missing)}. "
                f"Install the missing tools or pass skip_pdf=True."
            )
        summary_pdf = docx_to_pdf(summary_docx, toolchain_doctor_result=True)
        change_log_pdf = docx_to_pdf(change_log_docx, toolchain_doctor_result=True)

    changes_csv = build_changes_csv(changes)
    changes_json = build_changes_json(changes)
    readme = build_readme(export_id, snapshot_id, engagement_id, export_time, skip_pdf)

    # Build ZIP
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:

        def _add(name: str, data: BytesIO | StringIO) -> str:
            arcname = f"{dirname}/{name}"
            content = data.getvalue()
            if isinstance(content, str):
                content = content.encode("utf-8")
            zf.writestr(arcname, content)
            return sha256(content).hexdigest()

        checksums: dict[str, str] = {}
        checksums["Summary.xlsx"] = _add("Summary.xlsx", summary_xlsx)
        checksums["Change-Log.xlsx"] = _add("Change-Log.xlsx", change_log_xlsx)
        checksums["Summary.docx"] = _add("Summary.docx", summary_docx)
        checksums["Change-Log.docx"] = _add("Change-Log.docx", change_log_docx)
        if summary_pdf is not None:
            checksums["Summary.pdf"] = _add("Summary.pdf", summary_pdf)
        if change_log_pdf is not None:
            checksums["Change-Log.pdf"] = _add("Change-Log.pdf", change_log_pdf)
        checksums["Changes.csv"] = _add("Changes.csv", changes_csv)
        checksums["Changes.json"] = _add("Changes.json", changes_json)
        checksums["Read-Me.txt"] = _add("Read-Me.txt", readme)

        manifest_buf = build_manifest(
            export_id, snapshot_id, engagement_id,
            checksums, len(changes), len(warnings),
        )
        _add("manifest.json", manifest_buf)

    zip_buf.seek(0)
    return zip_buf
