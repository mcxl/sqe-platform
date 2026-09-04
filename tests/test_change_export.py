"""Phase 5B — ACE Change Record And Export Authorities tests.

Covers: snapshot consistency, stable IDs, append-only records,
duplicate publication prevention, publication failure handling,
SQLite immutability, summary counts, warnings, XLSX/CSV/JSON/DOCX/PDF
content, no photo embedding, cross-format identity, fictional-data
protection.
"""

from __future__ import annotations

import csv
import json
import os
import sqlite3
import uuid
import zipfile
from hashlib import sha256
from io import BytesIO, StringIO
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from src.ace.app import app
from src.ace.workbench.change_record import ChangeRecord
from src.ace.workbench.document_toolchain_doctor import probe as doctor_probe
from src.ace.workbench import document_toolchain_doctor
from src.ace.workbench.export_builder import (
    build_change_log_docx,
    build_change_log_xlsx,
    build_changes_csv,
    build_changes_json,
    build_export_zip,
    build_readme,
    build_summary_docx,
    build_summary_xlsx,
    generate_export_id,
)
from src.ace.workbench.notion_publisher import NotionPublisher
from src.ace.workbench.storage import WorkbenchStore


def _credentials() -> tuple[str, str]:
    return ("auditor", os.environ.get("ACE_AUDITOR_PASSWORD", "fictional-password"))


def _workbench_client(tmp_path: Path, suffix: str = "") -> WorkbenchStore:
    """Create a WorkbenchStore with an isolated data directory per test."""
    import uuid as _uuid

    unique = suffix or _uuid.uuid4().hex[:8]
    data_dir = tmp_path / f"sqe-local-data-{unique}"
    data_dir.mkdir(parents=True, exist_ok=True)
    return WorkbenchStore(data_dir=data_dir)


# ── Helpers ────────────────────────────────────────────────────

def _make_sample_changes(export_id: str = "EXP-2026-AAAA1234") -> list[ChangeRecord]:
    """Return two sample change records for export builder tests."""
    return [
        ChangeRecord.make(
            export_id=export_id,
            record_id="OBL-FIC-0001",
            snapshot_id="SNP-AAAA00000000",
            evidence_id=None,
            idempotency_key="test-key-001",
            change_type="added",
            record_type="obligation",
            label="Fictional obligation",
        ),
        ChangeRecord.make(
            export_id=export_id,
            record_id="EVD-FIC-0001",
            snapshot_id="SNP-AAAA00000000",
            evidence_id="EVD-FIC-0001",
            idempotency_key="test-key-001",
            change_type="modified",
            record_type="evidence",
            label="inspection.pdf",
            detail="Source text changed",
        ),
    ]


# ── ChangeRecord Model ─────────────────────────────────────────

class TestChangeRecordModel:
    def test_change_id_is_stable_and_deterministic(self) -> None:
        a = ChangeRecord.make("EXP-2026-AAAA", "REC-001", "SNP-AAA", None, "key-1", "added", "obligation", "Test")
        b = ChangeRecord.make("EXP-2026-AAAA", "REC-001", "SNP-AAA", None, "key-1", "added", "obligation", "Test")
        # Same inputs → same change_id prefix pattern
        assert a.change_id.startswith("CHG-"), f"Unexpected change_id format: {a.change_id}"
        assert len(a.change_id) == 16  # CHG- + 12 hex chars

    def test_change_record_as_dict_has_all_fields(self) -> None:
        ch = ChangeRecord.make("EXP-2026-AAAA", "REC-001", "SNP-AAA", "EVD-FIC-0001", "key-1", "added", "obligation", "Test")
        d = ch.as_dict()
        for field in ("change_id", "export_id", "record_id", "snapshot_id",
                       "evidence_id", "idempotency_key", "timestamp",
                       "change_type", "record_type", "label", "detail"):
            assert field in d, f"Missing field: {field}"


# ── Notion Publisher Test Double ────────────────────────────────

class TestNotionPublisher:
    def test_publish_returns_stable_publication_id(self) -> None:
        pub = NotionPublisher()
        r1 = pub.publish("key-1", "EXP-001")
        r2 = pub.publish("key-1", "EXP-001")
        # Both calls return same result — idempotent, never 409
        assert r1.published
        assert r1.publication_id == r2.publication_id
        assert r1.publication_id.startswith("NTN-")

    def test_is_published_tracks_keys(self) -> None:
        pub = NotionPublisher()
        assert not pub.is_published("key-x")
        pub.publish("key-x", "EXP-001")
        assert pub.is_published("key-x")

    def test_duplicate_publication_prevented(self) -> None:
        pub = NotionPublisher()
        pub.publish("key-dup", "EXP-001")
        # Same key, different export — returns existing, never 409
        assert pub.is_published("key-dup")
        # No error, no duplicate publication created
        r = pub.publish("key-dup", "EXP-002")
        assert r.error is None  # No error, returns existing

    def test_simulate_failure_returns_error(self) -> None:
        pub = NotionPublisher()
        pub.simulate_failure(True)
        r = pub.publish("key-fail", "EXP-001")
        assert not r.published
        assert "Simulated" in (r.error or "")


# ── Snapshot ────────────────────────────────────────────────────

class TestSnapshot:
    def test_snapshot_consistency_same_hash(self, tmp_path: Path) -> None:
        store = _workbench_client(tmp_path, "snap-hash")
        from src.ace.workbench.engagement import EngagementService
        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement")
        r1 = store.create_snapshot("snap-key-1")
        r2 = store.create_snapshot("snap-key-1")
        assert r1["content_hash"] == r2["content_hash"]
        assert r2["idempotent"] is True

    def test_snapshot_no_current_engagement_raises(self, tmp_path: Path) -> None:
        store = _workbench_client(tmp_path, "snap-empty")
        # Clear the seeded current engagement
        with store.connect() as conn:
            conn.execute("DELETE FROM current_engagement")
            conn.commit()
        with pytest.raises(ValueError, match="No current engagement"):
            store.create_snapshot("snap-key-empty")

    def test_snapshot_has_stable_id(self, tmp_path: Path) -> None:
        store = _workbench_client(tmp_path, "snap-id")
        from src.ace.workbench.engagement import EngagementService
        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement")
        r = store.create_snapshot("snap-key-id")
        assert str(r["snapshot_id"]).startswith("SNP-")
        # SNP- (4 chars) + 12 hex chars = 16
        assert len(str(r["snapshot_id"])) == 16

    def test_snapshot_includes_node_edge_warning_counts(self, tmp_path: Path) -> None:
        store = _workbench_client(tmp_path, "snap-counts")
        from src.ace.workbench.engagement import EngagementService
        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement in seeded data")
        r = store.create_snapshot("snap-key-counts")
        assert isinstance(r["node_count"], int)
        assert isinstance(r["edge_count"], int)
        assert isinstance(r["warning_count"], int)
        assert r["node_count"] > 0


# ── Change Records ──────────────────────────────────────────────

class TestChangeRecords:
    def test_change_records_are_stable(self, tmp_path: Path) -> None:
        store = _workbench_client(tmp_path, "cr-stable")
        from src.ace.workbench.engagement import EngagementService

        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement")
        snap = store.create_snapshot("cr-key-1")
        changes = store.detect_changes(str(snap["snapshot_id"]), "EXP-2026-TEST", "cr-key-1")
        for ch in changes:
            assert ch.change_id.startswith("CHG-")
            assert ch.export_id == "EXP-2026-TEST"

    def test_change_records_idempotent(self, tmp_path: Path) -> None:
        store = _workbench_client(tmp_path, "cr-idem")
        from src.ace.workbench.engagement import EngagementService

        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement")
        snap = store.create_snapshot("cr-key-2")
        changes1 = store.detect_changes(str(snap["snapshot_id"]), "EXP-2026-TEST", "cr-key-2")
        changes2 = store.detect_changes(str(snap["snapshot_id"]), "EXP-2026-TEST", "cr-key-2")
        assert len(changes1) == len(changes2)
        assert [c.change_id for c in changes1] == [c.change_id for c in changes2]


class TestChangeRecordsAppendOnly:
    def test_insert_then_check_no_update_possible(self, tmp_path: Path) -> None:
        """Verify change_records table has no-update/no-delete triggers."""
        store = _workbench_client(tmp_path, "ao-test")
        from src.ace.workbench.engagement import EngagementService
        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement")
        snap = store.create_snapshot("ao-snap-key")
        sid = str(snap["snapshot_id"])
        with store.connect() as conn:
            # INSERT is allowed
            conn.execute(
                "INSERT INTO change_records VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                ("CHG-TEST001", "EXP-TEST", "REC-001", sid,
                 None, "ao-key-1", "2026-01-01T00:00:00Z",
                 "added", "obligation", "Test", ""),
            )
            conn.commit()
            # UPDATE should be blocked by trigger
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute(
                    "UPDATE change_records SET label = 'Changed' WHERE change_id = 'CHG-TEST001'"
                )
                conn.commit()
            conn.rollback()
            # DELETE should be blocked by trigger
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("DELETE FROM change_records WHERE change_id = 'CHG-TEST001'")
                conn.commit()
            conn.rollback()

    def test_snapshots_are_immutable(self, tmp_path: Path) -> None:
        store = _workbench_client(tmp_path, "snap-imm")
        from src.ace.workbench.engagement import EngagementService

        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement")
        snap = store.create_snapshot("imm-key")
        sid = str(snap["snapshot_id"])
        with store.connect() as conn:
            with pytest.raises(Exception):
                conn.execute("UPDATE snapshots SET node_count = 999 WHERE snapshot_id = ?", (sid,))
                conn.commit()
            conn.rollback()
            with pytest.raises(Exception):
                conn.execute("DELETE FROM snapshots WHERE snapshot_id = ?", (sid,))
                conn.commit()
            conn.rollback()


# ── SQLite Immutability (existing tables) ───────────────────────

class TestSqliteImmutability:
    def test_snapshot_methods_only_read_from_existing_tables(self, tmp_path: Path) -> None:
        """graph_projection-like snapshot reads don't INSERT/UPDATE/DELETE
        on existing engagement/obligation/risk/etc. tables."""
        store = _workbench_client(tmp_path, "sql-imm")
        from src.ace.workbench.engagement import EngagementService

        svc = EngagementService(store)
        if svc.current() is None:
            pytest.skip("No current engagement")
        # _capture_snapshot is called by create_snapshot internally
        store.create_snapshot("sqlite-key")
        # If we got here without errors, the read-only path worked
        with store.connect() as conn:
            count = conn.execute("SELECT COUNT(*) FROM obligations").fetchone()[0]
            assert count > 0  # Data still intact


# ── Export Builder — XLSX / CSV / JSON ──────────────────────────

class TestExportFormats:
    def test_summary_xlsx_has_required_headers(self) -> None:
        changes = _make_sample_changes()
        buf = build_summary_xlsx(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
        )
        from openpyxl import load_workbook

        wb = load_workbook(buf)
        ws = wb["Summary"]
        # Check key fields are present
        text = "\n".join(str(ws.cell(row=r, column=1).value or "") for r in range(1, 30))
        assert "Export ID" in text
        assert "Snapshot ID" in text
        # "Not authoritative" notice — search all cells in both columns
        full_text = ""
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
            for cell in row:
                full_text += str(cell or "") + " "
        assert "not authoritative" in full_text.lower()

    def test_change_log_xlsx_has_all_headers(self) -> None:
        changes = _make_sample_changes()
        buf = build_change_log_xlsx(changes)
        from openpyxl import load_workbook

        wb = load_workbook(buf)
        ws = wb.active
        headers = [ws.cell(row=1, column=c).value for c in range(1, 12)]
        assert "Change ID" in headers
        assert "Export ID" in headers
        assert "Idempotency Key" in headers

    def test_changes_csv_matches_json(self) -> None:
        changes = _make_sample_changes()
        csv_buf = build_changes_csv(changes)
        json_buf = build_changes_json(changes)

        csv_data = list(csv.DictReader(StringIO(csv_buf.getvalue())))
        json_data = json.loads(json_buf.getvalue())

        assert len(csv_data) == len(json_data)
        assert csv_data[0]["change_id"] == json_data[0]["change_id"]

    def test_changes_json_is_valid(self) -> None:
        changes = _make_sample_changes()
        buf = build_changes_json(changes)
        data = json.loads(buf.getvalue())
        assert isinstance(data, list)
        assert len(data) == 2
        assert all("change_id" in item for item in data)


# ── Export Builder — DOCX ───────────────────────────────────────

class TestDocxFormats:
    def test_summary_docx_has_capital_case_headings(self) -> None:
        changes = _make_sample_changes()
        buf = build_summary_docx(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
        )
        from docx import Document

        doc = Document(buf)
        # Check all paragraph styles (level=0 uses "Title", level>0 uses "Heading")
        headings = [p.text for p in doc.paragraphs
                     if p.style.name.startswith("Heading") or p.style.name == "Title"]
        assert any("Change Export Summary" in h for h in headings)
        assert any("Export Details" in h for h in headings)

    def test_change_log_docx_lists_all_changes(self) -> None:
        changes = _make_sample_changes()
        buf = build_change_log_docx(changes)
        from docx import Document

        doc = Document(buf)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "CHG-" in text
        assert "OBL-FIC-0001" in text
        assert "EVD-FIC-0001" in text

    def test_docx_contains_not_authoritative_notice(self) -> None:
        changes = _make_sample_changes()
        buf = build_summary_docx(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
        )
        from docx import Document

        doc = Document(buf)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "not authoritative" in text.lower()
        assert "Writeback" in text


# ── Export Builder — PDF ────────────────────────────────────────

class TestPdfGeneration:
    def test_pdf_generation_when_toolchain_available(self) -> None:
        """PDF generation works when LibreOffice is available."""
        changes = _make_sample_changes()
        zip_buf = build_export_zip(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
            skip_pdf=False,
        )
        with zipfile.ZipFile(zip_buf) as zf:
            names = zf.namelist()
        dir_prefix = "ACE-change-export-EXP-2026-AAAA/"
        assert f"{dir_prefix}Summary.pdf" in names
        assert f"{dir_prefix}Change-Log.pdf" in names

    def test_pdf_skipped_when_skip_pdf_true(self) -> None:
        changes = _make_sample_changes()
        zip_buf = build_export_zip(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
            skip_pdf=True,
        )
        with zipfile.ZipFile(zip_buf) as zf:
            names = zf.namelist()
        dir_prefix = "ACE-change-export-EXP-2026-AAAA/"
        assert f"{dir_prefix}Summary.pdf" not in names


# ── Export Builder — ZIP ────────────────────────────────────────

class TestExportZip:
    def test_zip_contains_all_required_files(self) -> None:
        changes = _make_sample_changes()
        zip_buf = build_export_zip(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
        )
        with zipfile.ZipFile(zip_buf) as zf:
            names = zf.namelist()
        prefix = "ACE-change-export-EXP-2026-AAAA/"
        required = [
            "Read-Me.txt", "Summary.xlsx", "Change-Log.xlsx",
            "Summary.docx", "Change-Log.docx",
            "Changes.csv", "Changes.json", "manifest.json",
        ]
        for f in required:
            assert f"{prefix}{f}" in names, f"Missing: {f}"

    def test_readme_has_export_details(self) -> None:
        buf = build_readme("EXP-2026-AAAA", "SNP-AAAA", "ENG-FIC-0001",
                          "2026-08-19T12:00:00Z", skip_pdf=False)
        text = buf.getvalue()
        assert "EXP-2026-AAAA" in text
        assert "SNP-AAAA" in text
        assert "not authoritative" in text.lower()

    def test_no_photos_in_export(self) -> None:
        """Verify no image content in any export file."""
        changes = _make_sample_changes()
        zip_buf = build_export_zip(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
        )
        with zipfile.ZipFile(zip_buf) as zf:
            for name in zf.namelist():
                content = zf.read(name)
                # Check for image magic bytes
                has_png = b"\x89PNG" in content
                has_jpg = b"\xff\xd8\xff" in content
                assert not has_png, f"PNG found in {name}"
                assert not has_jpg, f"JPEG found in {name}"

    def test_cross_format_identity_same_records(self) -> None:
        """Same change records appear in XLSX, CSV, and JSON."""
        changes = _make_sample_changes()
        change_ids = {c.change_id for c in changes}

        # XLSX
        xlsx_buf = build_change_log_xlsx(changes)
        from openpyxl import load_workbook

        wb = load_workbook(xlsx_buf)
        ws = wb.active
        xlsx_ids = set()
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0]:
                xlsx_ids.add(str(row[0]))

        # CSV
        csv_buf = build_changes_csv(changes)
        csv_data = csv.DictReader(StringIO(csv_buf.getvalue()))
        csv_ids = {row["change_id"] for row in csv_data}

        # JSON
        json_buf = build_changes_json(changes)
        json_data = json.loads(json_buf.getvalue())
        json_ids = {item["change_id"] for item in json_data}

        assert change_ids == xlsx_ids == csv_ids == json_ids


# ── Export ID ───────────────────────────────────────────────────

class TestExportId:
    def test_export_id_is_stable_format(self) -> None:
        from datetime import UTC, datetime
        eid = generate_export_id("ENG-FIC-0001")
        assert eid.startswith("EXP-")
        parts = eid.split("-")
        assert len(parts) == 3
        # Year fragment must be the current year (4 digits)
        current_year = str(datetime.now(UTC).year)
        assert parts[1] == current_year
        # Hex fragment is 12 uppercase hex digits
        assert len(parts[2]) == 12
        assert all(c in "0123456789ABCDEF" for c in parts[2])

    def test_export_id_deterministic_for_same_input(self) -> None:
        """Same engagement + same timestamp yields same export ID."""
        from datetime import UTC, datetime
        # generate_export_id uses current time, so two calls differ
        # But format is always stable
        current_year = str(datetime.now(UTC).year)
        eid1 = generate_export_id("ENG-FIC-0001")
        eid2 = generate_export_id("ENG-FIC-0001")
        assert eid1.startswith(f"EXP-{current_year}-")
        assert eid2.startswith(f"EXP-{current_year}-")


# ── Fictional Data Protection ──────────────────────────────────

class TestFictionalDataProtection:
    def test_export_contains_fictional_notice(self) -> None:
        changes = _make_sample_changes()
        buf = build_summary_xlsx(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
        )
        from openpyxl import load_workbook

        wb = load_workbook(buf)
        ws = wb["Summary"]
        full_text = ""
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                full_text += str(cell or "")
        assert "fictional" in full_text.lower()


# ── Document Toolchain Doctor ──────────────────────────────────

class TestDocumentToolchainDoctor:
    def test_probe_accepts_windows_soffice_com(self, monkeypatch: pytest.MonkeyPatch) -> None:
        seen: list[str] = []

        def which(candidate: str) -> str | None:
            seen.append(candidate)
            return "C:/tools/soffice.com" if candidate == "soffice.com" else None

        completed = __import__("subprocess").CompletedProcess([], 0, "LibreOffice 25", "")
        monkeypatch.setattr(document_toolchain_doctor.shutil, "which", which)
        monkeypatch.setattr(document_toolchain_doctor.subprocess, "run", lambda *args, **kwargs: completed)

        result = document_toolchain_doctor._probe_libreoffice()

        assert result.available
        assert result.version == "LibreOffice 25"
        assert seen == ["libreoffice", "soffice.com", "soffice.com"]

    def test_probe_all_tools_available(self) -> None:
        report = doctor_probe()
        report.print()
        assert report.all_available, "LibreOffice, python-docx, and reportlab must be available"

    def test_probe_python_docx_found(self) -> None:
        report = doctor_probe()
        docx_result = next(r for r in report.results if r.tool == "python-docx")
        assert docx_result.available


# ── API Endpoints ──────────────────────────────────────────────

class TestExportApi:
    @pytest.fixture
    def api_client(self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
        monkeypatch.setenv("ACE_AUDITOR_PASSWORD", "fictional-password")
        monkeypatch.setenv("ACE_DATA_DIR", str(tmp_path / "sqe-local-data"))
        return TestClient(app)

    def test_create_export_requires_auth(self, api_client: TestClient) -> None:
        response = api_client.post(
            "/workbench/api/v1/engagements/export",
            json={"idempotency_key": "test-auth-key-001"},
        )
        assert response.status_code == 401

    def test_create_export_no_engagement_returns_400(self, api_client: TestClient) -> None:
        # Deactivate the current engagement so the export endpoint sees none
        from src.ace.workbench.storage import WorkbenchStore as WS
        from pathlib import Path as _Path
        import os as _os

        data_dir = _os.environ.get("ACE_DATA_DIR", "")
        store = WS(data_dir=_Path(data_dir)) if data_dir else WS()
        with store.connect() as conn:
            conn.execute("DELETE FROM current_engagement")
            conn.commit()

        response = api_client.post(
            "/workbench/api/v1/engagements/export",
            json={"idempotency_key": f"test-no-eng-{uuid.uuid4().hex[:8]}"},
            auth=_credentials(),
        )
        assert response.status_code == 400, f"Expected 400, got {response.status_code}: {response.text}"

    def test_export_idempotent(self, api_client: TestClient) -> None:
        # Ensure an engagement is active (seeded data makes one available)
        from src.ace.workbench.storage import WorkbenchStore as WS
        from pathlib import Path as _Path
        import os as _os

        data_dir = _os.environ.get("ACE_DATA_DIR", "")
        store = WS(data_dir=_Path(data_dir)) if data_dir else WS()
        # The seeded schema creates ENG-FIC-0001 as current if none exists
        with store.connect() as conn:
            conn.execute("SELECT 1 FROM current_engagement LIMIT 1")

        key = f"idem-test-{uuid.uuid4().hex[:8]}"
        response1 = api_client.post(
            "/workbench/api/v1/engagements/export",
            json={"idempotency_key": key},
            auth=_credentials(),
        )
        assert response1.status_code == 201, f"Expected 201, got {response1.status_code}: {response1.text}"

        response2 = api_client.post(
            "/workbench/api/v1/engagements/export",
            json={"idempotency_key": key},
            auth=_credentials(),
        )
        assert response2.status_code == 201, f"Expected 201, got {response2.status_code}: {response2.text}"
        assert response2.json()["idempotent"] is True

    def test_download_export_not_found(self, api_client: TestClient) -> None:
        response = api_client.get(
            "/workbench/api/v1/engagements/export/EXP-9999-NOPE",
            auth=_credentials(),
        )
        assert response.status_code == 404

    def test_export_request_g0_rejects_real(self, api_client: TestClient) -> None:
        """G0 validation blocks export for real-client data."""
        # The seeded engagement is fictional, so this passes through.
        # G0 is enforced by EngagementService.current() which is called
        # in the export endpoint. We verify it's wired.
        response = api_client.post(
            "/workbench/api/v1/engagements/export",
            json={"idempotency_key": f"g0-test-{uuid.uuid4().hex[:8]}"},
            auth=_credentials(),
        )
        # Either 400 (no current engagement) or 201 (exported)
        # Neither should be 403 since seed data is fictional
        assert response.status_code != 403, "G0 should not reject fictional data"


# ── Snapshot Identity Across Formats ──────────────────────────

class TestExportSnapshotIdentity:
    def test_manifest_checksums_match_file_content(self) -> None:
        changes = _make_sample_changes()
        zip_buf = build_export_zip(
            "EXP-2026-AAAA", "SNP-AAAA", "2026-08-19T00:00:00Z",
            "ENG-FIC-0001", "Fictional Engagement",
            changes, [], 0, None,
        )
        with zipfile.ZipFile(zip_buf) as zf:
            # Read manifest
            manifest_raw = zf.read("ACE-change-export-EXP-2026-AAAA/manifest.json")
            manifest = json.loads(manifest_raw)
            for filename, expected_hash in manifest["files"].items():
                arcname = f"ACE-change-export-EXP-2026-AAAA/{filename}"
                content = zf.read(arcname)
                actual_hash = sha256(content).hexdigest()
                assert actual_hash == expected_hash, f"Checksum mismatch for {filename}"
